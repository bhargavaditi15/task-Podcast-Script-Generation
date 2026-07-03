"""Extract plain text from an uploaded PDF / DOCX / TXT file.

All failures (corrupt file, password-protected PDF, unsupported legacy .doc,
unreadable encoding) are surfaced as DocumentParseError with a message that's
safe to show directly in the Streamlit UI or CLI output.
"""

import io
import pathlib


class DocumentParseError(Exception):
    """User-facing document parsing failure."""


def extract_text(file_bytes: bytes, filename: str) -> str:
    # Choose the parser based on file extension and return plain text.
    if not file_bytes:
        raise DocumentParseError(f"'{filename}' is empty (0 bytes).")

    ext = pathlib.Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)
    if ext == ".docx":
        return _extract_docx(file_bytes, filename)
    if ext == ".doc":
        raise DocumentParseError(
            f"'{filename}' is a legacy .doc file, which this tool can't parse reliably. "
            "Please re-save it as .docx or .pdf and re-upload."
        )
    if ext == ".txt":
        return _extract_txt(file_bytes, filename)

    raise DocumentParseError(f"Unsupported file type '{ext or '(none)'}' for '{filename}'. Supported: PDF, DOC/DOCX, TXT.")


def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise DocumentParseError("The 'pdfplumber' package is not installed. Run: pip install pdfplumber") from exc

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:  # noqa: BLE001 - pdfplumber/pdfminer raise assorted types
        msg = str(exc).lower()
        if "password" in msg or "encrypt" in msg:
            raise DocumentParseError(f"'{filename}' appears to be password-protected. Please upload an unlocked copy.") from exc
        raise DocumentParseError(f"'{filename}' could not be read as a PDF (it may be corrupted): {exc}") from exc

    text = "\n\n".join(p for p in pages if p.strip())
    if not text.strip():
        raise DocumentParseError(
            f"'{filename}' has no extractable text (it may be a scanned/image-only PDF). "
            "Please upload a text-based document instead."
        )
    return text


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    try:
        import docx
    except ImportError as exc:
        raise DocumentParseError("The 'python-docx' package is not installed. Run: pip install python-docx") from exc

    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.extend(cell.text for cell in row.cells)
    except Exception as exc:  # noqa: BLE001
        raise DocumentParseError(f"'{filename}' could not be read as a DOCX file (it may be corrupted): {exc}") from exc

    text = "\n".join(p for p in paragraphs if p and p.strip())
    if not text.strip():
        raise DocumentParseError(f"'{filename}' has no extractable text content.")
    return text


def _extract_txt(file_bytes: bytes, filename: str) -> str:
    encoding = "utf-8"
    try:
        import chardet

        detected = chardet.detect(file_bytes)
        if detected and detected.get("encoding") and detected.get("confidence", 0) > 0.5:
            encoding = detected["encoding"]
    except ImportError:
        pass

    try:
        text = file_bytes.decode(encoding, errors="strict")
    except (UnicodeDecodeError, LookupError):
        text = file_bytes.decode("utf-8", errors="replace")

    if not text.strip():
        raise DocumentParseError(f"'{filename}' is empty or contains no readable text.")
    return text
